import dateparser
from docx import Document
from dotenv import load_dotenv
import logging
import os
import pytz
import re

logger = logging.getLogger(__name__)

load_dotenv()

PRD_PATH = os.getenv('LOCAL_PRDs')


def format_datetime_utc(dt):
    if dt.tzinfo is None:
        dt = dt.replace(tzinfo=pytz.UTC)
    else:
        dt = dt.astimezone(pytz.UTC)
    return dt.strftime('%Y-%m-%dT%H:%M:%S.%f')[:-3] + 'Z'


def get_prds(path=PRD_PATH):
    files = {}
    logger.info(f'Fetching PRDs from {PRD_PATH}')
    for f in os.listdir(path):
        file_path = os.path.join(PRD_PATH, f)
        if os.path.isfile(file_path):
            files[f] = file_path
    return files


def extract_date_from_document(doc_path):
    parsed_dates = []
    doc = Document(doc_path)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parsed = None
                if len(cell.text) > 8 or re.fullmatch(cell.text, '/d+(/./d+)+'):
                    parsed = dateparser.parse(cell.text, settings={'TIMEZONE': 'UTC'})   
                if parsed:
                    parsed_dates.append(parsed)
        if parsed_dates:
            break
    if parsed_dates:
        created_at = format_datetime_utc(min(parsed_dates))
        updated_at = format_datetime_utc(max(parsed_dates))
        return created_at, updated_at
    return (None, None)


def extract_text_from_docx(path):
    doc = Document(path)
    return "\n".join([para.text for para in doc.paragraphs if para.text.strip()])