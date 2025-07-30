from docx import Document
import logging
import os
from utils.db_utils import get_similar_docs
from utils.ollama_utils import get_embedding, generate_prd
from utils.text_utils import chunk_text

GENERATED_PRD_PATH = os.getenv('GENERATED_PRD_PATH')

logger = logging.getLogger(__name__)

def process_input_text(text):
    if not text:
        logger.debug('No Text to process')
        return
    
    chunks = chunk_text(text)
    logger.info(f'{len(chunks)} chunks created for Input BR')
    context_docs = {}
    for i, chunk in enumerate(chunks):
        logger.info(f'Embedding chunk {i} in Input Text')
        embedding = get_embedding(chunk)
        results = get_similar_docs(embedding)
        for k,v in results.items():
            if k not in context_docs:
                context_docs[k] = v
    context_str = '\n\n'.join(str(v) for v in context_docs.values())
    response = generate_prd(text, context_str)
    write_prd_to_file(response)


def write_prd_to_file(prd):
    doc = Document()
    for line in prd.splitlines():
        line = line.strip()
        if not line:
            doc.add_paragraph("")
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith("1.") or line.startswith("2.") or line.startswith("3."):
            doc.add_paragraph(line)
        else:
            doc.add_paragraph(line)  
    save_path = f'{GENERATED_PRD_PATH}/prd.docx'
    doc.save(save_path)

